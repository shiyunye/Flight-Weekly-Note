from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.table import WD_ROW_HEIGHT


def set_table_style(table, border_color="000000", border_width=4, fill_color="FFFFFF"):

    table_properties = table._element.find(qn("w:tblPr"))
    if table_properties is None:
        table_properties = OxmlElement("w:tblPr")
        table._element.insert(0, table_properties)

    # Set table borders
    tbl_borders = OxmlElement("w:tblBorders")
    for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_element = OxmlElement(f"w:{border}")
        border_element.set(qn("w:val"), "single")
        border_element.set(qn("w:sz"), str(border_width))  # Border width in eighths of a point
        border_element.set(qn("w:space"), "0")
        border_element.set(qn("w:color"), border_color)
        tbl_borders.append(border_element)
    table_properties.append(tbl_borders)

    # Set cell background color
    for row in table.rows:
        for cell in row.cells:
            # Retrieve or create the <tcPr> (table cell properties)
            cell_properties = cell._element.find(qn("w:tcPr"))
            if cell_properties is None:
                cell_properties = OxmlElement("w:tcPr")
                cell._element.insert(0, cell_properties)

            # Set cell shading (background color)
            cell_shading = OxmlElement("w:shd")
            cell_shading.set(qn("w:val"), "clear")
            cell_shading.set(qn("w:fill"), fill_color)  # Background color
            cell_properties.append(cell_shading)
            
def shade_cell(cell, color):
    """Apply background color to a cell."""
    cell_properties = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell_properties.append(shading)

def set_table_width(table, width_pct='2500'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), width_pct)
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

def adjust_cell_dimensions(table):
    if table is None:
        raise ValueError("The table is None. Make sure you created it correctly before adjusting dimensions.")
    
    table.allow_autofit = False  # Disable autofit to manually control dimensions

    for row in table.rows:
        row.height = Inches(0.4)  # Set row height

        for cell in row.cells:
            cell.width = Inches(0.8)  # Set cell width

def index_width_dimensions(table):
    """
    Adjusts the dimensions of a table, including row height and column widths.
    """
    # Disable autofit to allow manual size adjustment
    table.allow_autofit = False

    # Set uniform row height
    for row in table.rows:
        row.height = Inches(0.4)

    # # Adjust column width based on content length
    # for col_idx, column in enumerate(table.columns):
    #     max_text_length = max(
    #         (len(cell.text.strip()) for cell in column.cells), default=0
    #     )  # Get the max text length in the column
    #     # Estimate column width based on the longest text (adjust scaling factor if needed)
    #     column_width = max(Inches(1), Inches(max_text_length * 0.1))
    #     column.width = column_width

def center_text_in_cell(cell):
    """Center text horizontally and vertically in a table cell."""
    # Horizontal alignment
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Vertical alignment
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def center_bottom_text_in_cell(cell):
    """Center text horizontally and vertically in a table cell."""
    # Horizontal alignment
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Vertical alignment
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'bottom')
    tcPr.append(vAlign)

def set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=False, color=None):

    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
     
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)

def bold_highlight_total_cells(table):
    rows_to_highlight = set()
    cols_to_highlight = set()

    # Check the header row for 'Total' in column headers
    header_row = table.rows[0]
    for col_idx, cell in enumerate(header_row.cells):
        if "total" in cell.text.lower().strip():
            cols_to_highlight.add(col_idx)

    # Check all rows for 'Total' in the index (first cell of each row)
    for row_idx, row in enumerate(table.rows):
        if "total" in row.cells[0].text.lower().strip():
            rows_to_highlight.add(row_idx)

    # Apply formatting to the identified rows and columns
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if row_idx in rows_to_highlight or col_idx in cols_to_highlight:
                shade_cell(cell, 'D3D3D3')  # Custom shading function
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True  # Bold the text

def remove_cell_borders(cell):
    # Get or add the cell properties (tcPr)
    tc_pr = cell._element.get_or_add_tcPr()

    # Set all borders (top, bottom, left, right) to white
    for border in ["top", "bottom", "left", "right"]:
        border_element = tc_pr.find(qn(f"w:{border}"))
        if border_element is None:
            border_element = OxmlElement(f"w:{border}")
            tc_pr.append(border_element)
        
        border_element.set(qn("w:val"), "nil") 
        border_element.set(qn("w:color"), "FFFFFF")            

def format_percentage_values_other(table):

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    value = run.text.strip()  # Ensure text is not empty
                    if not value:
                        continue
                    
                    font_color = RGBColor(0, 0, 0)  # Default color is black

                    try:
                        # Check if the value is a percentage
                        if value.endswith('%'):
                            num = float(value.rstrip('%'))  # Convert to float after stripping %
                            run.text = f"{num:.1f}%"  # Format with one decimal place
                        else:
                            num = float(value)  # Convert to float
                            run.font.bold = True

                            if num >= 0:
                                run.text = f"+{num:.0f}"  # Add plus sign
                                font_color = RGBColor(106, 168, 79)  # Dark green
                            else: 
                                run.text = f"{num:.0f}"  # Keep negative sign
                                font_color = RGBColor(255, 0, 0)  # Red
                    except ValueError:
                        # Leave non-numerical text unchanged
                        pass

                    # Apply font styling
                    run.font.color.rgb = font_color
                    run.font.name = "Montserrat"

def format_percentage_values(
    table,
    header_row_idx=1,
    keywords=("Actual", "Net Tickets", "DAU", "Conversion", "ROI"),
    decimals=None,):
    # normalize keywords to lowercase
    keywords_lc = tuple(k.lower() for k in keywords)

    # mark columns that should stay black
    header_cells = table.rows[header_row_idx].cells
    col_is_key = []
    for c in header_cells:
        htxt = (c.text or "").strip().lower()
        col_is_key.append(any(k in htxt for k in keywords_lc))

    for r_idx, row in enumerate(table.rows):
        if r_idx == header_row_idx:
            continue  # skip header

        for c_idx, cell in enumerate(row.cells):
            is_key_col = col_is_key[c_idx] if c_idx < len(col_is_key) else False

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    value = (run.text or "").strip()
                    if not value.endswith('%'):
                        continue  # only handle percentages

                    raw = value[:-1].replace(',', '').strip()
                    if raw in ("", "+", "-", "+.", "-."):
                        continue

                    try:
                        num_value = float(raw)
                    except ValueError:
                        continue

                    # optional rounding
                    if decimals is not None:
                        num_value = round(num_value, decimals)

                    # always add +/- sign
                    if num_value > 0:
                        formatted_value = f"+{num_value}%"
                    elif num_value < 0:
                        formatted_value = f"{num_value}%"
                    else:
                        formatted_value = "0%"

                    # color rule:
                    #   if header contains keyword -> BLACK
                    #   else -> GREEN for +, RED for -, BLACK for 0
                    if is_key_col:
                        font_color = RGBColor(0, 0, 0)
                    else:
                        if num_value > 0:
                            font_color = RGBColor(106, 168, 79)  # green
                        elif num_value < 0:
                            font_color = RGBColor(255, 0, 0)    # red
                        else:
                            font_color = RGBColor(0, 0, 0)      # zero stays black

                    run.text = formatted_value
                    run.font.color.rgb = font_color
                    run.font.bold = True

def add_table_headers_index(table, title, column_titles, logos):
    
    # Table title in the first column
    
    header_cells = table.rows[0].cells
    remove_cell_borders(table.cell(0, 0))
    # print(title)
    table.cell(1, 0).text = title
    
    remove_cell_borders(table.cell(1, 0))
    table.cell(0, 0).merge(table.cell(1, 0))
    # center_text_in_cell(table.cell(1, 0))  # Center the text in the merged cell
    center_bottom_text_in_cell(table.cell(1, 0))  # Center the text in the merged cell

    for paragraph in header_cells[0].paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(14), bold=True)

 
    for cell in header_cells:
        cell.height = Inches(0.4) 


    # Add headerlogos with titles
    for header_logo in logos:
        logo_cell = table.cell(0, header_logo['start_col'])
        logo_cell.merge(table.cell(0, header_logo['end_col']))
        paragraph = logo_cell.paragraphs[0]

        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_text("\n\n")
        run.add_text("\n\n")
        # print(header_logo['path'])
        
        try:
            run.add_picture(header_logo['path'], width=Inches(0.4), height=Inches(0.4))
        except FileNotFoundError:
            paragraph.text = " "
            # print(header_logo)
        paragraph.add_run(f"\n{header_logo['title']}")
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
        if header_logo['title']== 'Total':
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True) 
            center_text_in_cell(cell)
        else:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
            center_text_in_cell(cell)

    # Add column titles
    for idx, col in enumerate(column_titles):
        cell = table.cell(1, idx + 1)
        col=col.split('_')[0]
        cell.text = col
        # print(col)
        center_text_in_cell(cell)
        for paragraph in cell.paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True)
        index_width_dimensions(table)

def populate_table_data(table, df):

    total_column_index = None
    total_column_indices = [col_index + 1  # Adjust for the table offset (logo column is index 0)
    for col_index, col_name in enumerate(df.columns)
    if "total" in col_name.lower()]

    for col_index, col_name in enumerate(df.columns):
        if col_name.lower() == 'total':
            total_column_index = col_index + 1

    for index, row_data in df.iterrows():
        cells = table.add_row().cells

        # Add Index logo to the first column
        index_logo_path = f'../Screenshots/weeklynote/{index}.png'
        # print(index)
        cell = cells[0]
        paragraph = cell.paragraphs[0]
    
        run = paragraph.add_run()
        
        if str(index).lower() == "total":  # Skip if the index is "Total"
            paragraph.add_run(f"{index}")
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10),bold=True)
            center_text_in_cell(cell)
        else:
            try:
                if "(" in str(index) and ")" in str(index):
                    # Define width and height as tuple (width, height)
                    size_dict = {
                        "American Airlines (AA)": (1.5, 0.21),
                        "Delta Air Lines (DL)": (1.4, 0.21),
                        "United Airlines (UA)": (1.4, 0.25),
                        "Southwest Airlines (WN)": (1.4, 0.21),
                        "Spirit Airlines (NK)": (1.0, 0.27),
                        "Frontier Airlines (F9)": (1.1, 0.24),
                        "Alaska Airlines (AS)": (0.8, 0.25),
                        "JetBlue Airways (B6)": (0.7, 0.25),
                        "Other": (0.2, 0.2)
                    }

                    # Get width and height or default to (1.0, 0.2)
                    logo_width, logo_height = size_dict.get(str(index), (1.0, 0.2))
                    run.add_picture(index_logo_path, height=Inches(float(logo_height)), width=Inches(float(logo_width)))
                    center_text_in_cell(cell)

                elif str(index)=='Other':
                        run.add_picture(index_logo_path, height=Inches(0.2),width=Inches(0.2))
                        # set_font(paragraph, font_name="Montserrat", font_size=Pt(7),bold=False)
                        paragraph.add_run(f"\n{index}")
                        set_font(paragraph, font_name="Montserrat", font_size=Pt(7),bold=False)
                else:
                    run.add_picture(index_logo_path, height=Inches(0.4),width=Inches(0.4))
                    set_font(paragraph, font_name="Montserrat", font_size=Pt(7),bold=False)
                    paragraph.add_run(f"\n{index}")
                    # print(index)
                    set_font(paragraph, font_name="Montserrat", font_size=Pt(7),bold=False)
                    center_text_in_cell(cell)
                
            except FileNotFoundError:
                    paragraph.text = "logo_path miss "

        center_text_in_cell(cell)


        # Populate data cells
        for col_idx, value in enumerate(row_data):
            cell = cells[col_idx + 1]
            # print(col_idx + 1)
            cell.width = Inches(5)
            # print(value)
            cell.text = str(value)
            center_text_in_cell(cell)
        
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name="Montserrat", font_size=Pt(10))

            if total_column_index and col_idx + 1 == total_column_index:
                shade_cell(cell, 'D3D3D3')

             # Highlight "Total" columns
            if col_idx + 1 in total_column_indices:
                shade_cell(cell, 'D3D3D3')

            # print("Row Indexes:", df.index.tolist())


        
        format_percentage_values(table)
                
   
        adjust_cell_dimensions(table)


        bold_highlight_total_cells(table)

def create_word_table(df, title, logos, word_document):

    # Define dynamic index column widths for different table titles
    index_column_widths = {
        "Business": Inches(1.146),
        "Carrier": Inches(1.646),
        "Channel": Inches(1),
        "Source": Inches(1),
        # "ROI": Inches(1)
    }

    # Set a default width if the title is not specified in the dictionary
    index_column_width = index_column_widths.get(title)
    
    # Create the table
    table = word_document.add_table(rows=2, cols=len(df.columns) + 1, style='Table Grid')
    
    # Apply table styling
    set_table_style(table, border_color="000000", border_width=8, fill_color="FFFFFF")
    
    # Adjust index column width for all rows
    table.columns[0].width = index_column_width


    # Add headers
    column_titles = df.columns.tolist()

    add_table_headers_index(table, title, column_titles, logos)

    # Populate data
    populate_table_data(table, df)

    if title == "Source":
        for cell_coords in [
            (4, 1), (4, 2), (4, 3),
        ]:
            remove_cell_borders(table.cell(*cell_coords))
    
    if title == "ROI":
        for i in range(2, 10):  # i goes from 2 to 9
            for j in range(1, 4):  # j goes from 1 to 3
                remove_cell_borders(table.cell(i, j))

            
    bold_highlight_total_cells(table)

    for row in table.rows:
        row.height = Inches(0.4)
        row.height_rule = WD_ROW_HEIGHT.EXACTLY

def clear_document(word_document):
    """Clear all content from a Word document."""
   
    for paragraph in word_document.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Remove all tables
    for table in word_document.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)

def create_summary_table(word_document,df_summary):

    # Add the index logo to the first column
    index_logo_path = '../Screenshots/weeklynote/priceline.png'

    # Create the table
    rows, cols = len(df_summary.index) + 2,len(df_summary.columns) + 2
    table = word_document.add_table(rows=rows, cols=cols, style='Table Grid')
    
    set_table_style(table, border_color="000000", border_width=8, fill_color="FFFFFF")
    
    # for cell_coords in [
    #     (0, 0), (5, 5),(5, 6), (5, 7), (5, 8), (5, 9),(6, 5),(6, 6), (6, 7), (6, 8), (6, 9),(7, 5),(7, 6), (7, 7), (7, 8), (7, 9),(8, 5),(8, 6), (8, 7), (8, 8), (8, 9),
    #     (9, 5),(9, 6), (9, 7), (9, 8), (9, 9)
    # ]:
    #     remove_cell_borders(table.cell(*cell_coords))

    # Add the logo to the first cell and merge header cells
    header_cell = table.cell(0, 0)
    header_paragraph = header_cell.paragraphs[0]
    header_paragraph.add_run().add_picture(index_logo_path, width=Inches(1.5), height=Inches(.52))
    center_text_in_cell(header_cell)

    table.cell(0, 2).merge(table.cell(0, 1)).merge(table.cell(0, 0))
    # table.cell(6,5).merge(table.cell(7, 9))
    table.cell(8,1).merge(table.cell(8, 0))
    table.cell(9,1).merge(table.cell(9, 0))


    # Define and populate column headers name
    columns = ['Metric', 'Net/Gross', 'Actual', 'Reporting Week', 'Previous Week', 'YTD','Reporting Week','Previous Week','YTD']
    for j, name in enumerate(columns):
        cell = table.cell(1, j)
        cell.text = name
        center_text_in_cell(cell)
        for paragraph in cell.paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True)


    # Add icons and labels for index col
    icons_and_labels = [
        ('../Screenshots/weeklynote/Tickets.png', 'Tickets', 2, 3),
        ('../Screenshots/weeklynote/Revenue.png', 'Revenue', 4, 5),
        ('../Screenshots/weeklynote/Normalized Tickets.png', 'Normalized Tickets', 6, 7)
    ]

    for icon_path, label, start_row, end_row in icons_and_labels:
        table.cell(end_row, 0).merge(table.cell(start_row, 0))
        cell = table.cell(start_row, 0)
        paragraph = cell.paragraphs[0]
        paragraph.add_run().add_picture(icon_path, width=Inches(0.4), height=Inches(0.4))
        paragraph.add_run('\n')
        paragraph.add_run(label)
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False)
        center_text_in_cell(cell)

    # add dau
    cell = table.cell(8,1).merge(table.cell(8, 0))
    paragraph = cell.paragraphs[0]
    paragraph.add_run().add_picture('../Screenshots/weeklynote/dau.png', width=Inches(0.4), height=Inches(0.4))
    paragraph.add_run("\n Daily Active Users")
    set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False)
    center_text_in_cell(cell)

    # add tsa
    cell = table.cell(9,1).merge(table.cell(9, 0))
    paragraph = cell.paragraphs[0]
    paragraph.add_run().add_picture('../Screenshots/weeklynote/tsa.png', width=Inches(1), height=Inches(0.33))
    paragraph.add_run("\n Market Share")
    set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False)
    center_text_in_cell(cell)

    # Add "vs. Last Year" and "vs. Plan" labels in the header row
    last_year_logo_path = '../Screenshots/weeklynote/lastyear.png'
    plan_logo_path = '../Screenshots/weeklynote/plan.png'
    
    # "vs. Last Year"
    cell = table.cell(0, 5).merge(table.cell(0, 3))
    paragraph = cell.paragraphs[0]
    paragraph.add_run().add_picture(last_year_logo_path, width=Inches(0.4), height=Inches(0.4))
    paragraph.add_run("\n vs. Last Year")
    set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
    center_text_in_cell(cell)

    # "vs. Plan"
    cell = table.cell(0, 8).merge(table.cell(0, 6))
    paragraph = cell.paragraphs[0]
    paragraph.add_run().add_picture(plan_logo_path, width=Inches(0.4), height=Inches(0.4))
    paragraph.add_run("\n vs. Plan")
    set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
    center_text_in_cell(cell)

    # Add "Net" and "Gross" alternately in col 2 for names
    for i in range(2, len(df_summary)):
        cell = table.cell(i, 1)
        cell.text = "Net" if (i - 2) % 2 == 0 else "Gross"
        paragraph = cell.paragraphs[0]  
        set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
        center_text_in_cell(cell)


    # Insert data from df_summary starting at cell (2,2)
    for i, row_data in enumerate(df_summary.itertuples(index=False), start=2):
        for j, value in enumerate(row_data, start=2):
            cell = table.cell(i, j)
            if (i, j) == (5, 2) or (i, j) == (4, 2):  
                cell.text = f"${value}"  # Add dollar sign in front of the value
                # print(value)
            else:
                cell.text = str(value)
                # print(value)
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name="Montserrat", font_size=Pt(10))
            center_text_in_cell(cell)
            format_percentage_values(table)
        format_percentage_values(table)

    set_table_width(table)

    adjust_cell_dimensions(table)

def create_others_table(word_document,df_others):
    
    # Define table dimensions
    rows, cols = (17, 9)
    table = word_document.add_table(rows=rows, cols=cols, style='Table Grid')
    set_table_style(table, border_color="000000", border_width=8, fill_color="FFFFFF")

        
    # Add header logos
    header_logo_path_1 = '../Screenshots/weeklynote/priceline.png'
    header_logo_path_2 = '../Screenshots/weeklynote/expedia.png'

    # # Remove unnecessary cell borders
    # for cell_coords in [
    #     (3, 6), (3, 7), (3, 8),(5, 6), (5, 7), (5, 8),(8, 6), (8, 7), (8, 8),(9, 6), (9, 7), (9, 8),
    #     (10, 6), (10, 7), (10, 8),(11, 6), (11, 7), (11, 8),(12, 6), (12, 7), (12, 8),(13, 6), (13, 7), (13, 8)
    # ]:
    #     remove_cell_borders(table.cell(*cell_coords))

    # Merge header cells and add logos
    table.cell(0, 0).merge(table.cell(1, 2))
    header_cell_1 = table.cell(0, 3).merge(table.cell(0, 5))
    # print(header_logo_path_1)
    logo_paragraph_1 = header_cell_1.paragraphs[0]
    logo_paragraph_1.add_run().add_picture(header_logo_path_1, width=Inches(1.5), height=Inches(0.52))
    center_text_in_cell(header_cell_1)

    header_cell_2 = table.cell(0, 6).merge(table.cell(0, 8))
    logo_paragraph_2 = header_cell_2.paragraphs[0]
    logo_paragraph_2.add_run().add_picture(header_logo_path_2, width=Inches(1.4), height=Inches(0.28))
    center_text_in_cell(header_cell_2)

    # Define and populate column headers
    columns = ['Actual', 'YoY (bps)', 'YoY PW (bps)', 'Actual', 'YoY (bps)', 'YoY PW (bps)']
    for j, name in enumerate(columns):
        cell = table.cell(1, j + 3)
        cell.text = name
        center_text_in_cell(cell)
        for paragraph in cell.paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True)


    # Add icons and labels for index columns
    icons_and_labels_col1 = [
        ('../Screenshots/weeklynote/Market Share.png', 'Market Share', 2, 8),
        ('../Screenshots/weeklynote/Parity.png', 'Parity', 9,11),
        ('../Screenshots/weeklynote/Share of Business.png', 'Share of Business', 12, 15),
    ]
    for icon_path, label, start_row, end_row in icons_and_labels_col1:
        table.cell(start_row, 0).merge(table.cell(end_row, 0))
        cell = table.cell(start_row, 0)
        paragraph = cell.paragraphs[0]
        paragraph.add_run().add_picture(icon_path, width=Inches(0.4), height=Inches(0.4))
        paragraph.add_run('\n'+label)
        # paragraph.add_run('\n')
        # paragraph.add_run('\n')
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False)
        center_text_in_cell(cell)
  
    icons_and_labels_col2 = [
        ('../Screenshots/weeklynote/arc.png', 'arc', 2, 5, 0.6, 0.25),
        ('../Screenshots/weeklynote/midt.png', 'MIDT', 6, 6, .4, .4),
        ('../Screenshots/weeklynote/sem.png', 'SEM Impressions', 7, 7, 0.4, 0.4),
        ('../Screenshots/weeklynote/expedia.png', 'Direct vs Expedia', 8, 8, 1, .2),
        ('../Screenshots/weeklynote/Kayak.png', 'Kayak Placement', 9, 9, 1.2, 0.23),
        ('../Screenshots/weeklynote/Skyscanner.png', 'Skyscanner Placement', 10, 10, 1.2, 0.21),
        ('../Screenshots/weeklynote/usoutbound.png', 'US Outbound', 11, 11, 0.4, 0.4),
        ('../Screenshots/weeklynote/Merchant.png', 'Merchant', 12, 13, 0.4, 0.4),
        ('../Screenshots/weeklynote/Deals.png', 'Deals', 14, 14, 0.4, 0.4),
        ('../Screenshots/weeklynote/Conversion.png', 'Conversion', 15, 15, 0.4, 0.4),
        ('../Screenshots/weeklynote/Bookability.png', 'Bookability', 16, 16, 0.4, 0.4),
    ]
    for icon_path, label, start_row, end_row, width, height in icons_and_labels_col2:
        if label=='Merchant':
            table.cell(start_row, 1).merge(table.cell(end_row, 1))
        elif label=='arc':
            table.cell(start_row, 1).merge(table.cell(end_row, 1))
        else:
            table.cell(start_row, 1).merge(table.cell(end_row, 2))

        cell = table.cell(start_row, 1)
        paragraph = cell.paragraphs[0]
        
        if label in {"arc", "tsa"}:
            paragraph.add_run().add_picture(icon_path, width=Inches(width), height=Inches(height))
        else:
            paragraph.add_run().add_picture(icon_path, width=Inches(width), height=Inches(height))
            paragraph.add_run('\n'+label)
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False)
        center_text_in_cell(cell)

   # Add "Retail" text and set its font
    retail_cell = table.cell(12, 2)
    retail_paragraph = retail_cell.paragraphs[0]
    retail_paragraph.text = "Retail"
    set_font(retail_paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
    center_text_in_cell(retail_cell)

    # Add "Total" text and set its font
    total_cell = table.cell(13, 2)
    total_paragraph = total_cell.paragraphs[0]
    total_paragraph.text = "Total"
    set_font(total_paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
    center_text_in_cell(total_cell)

     # Add "ARC" text and set its font
    retail_cell = table.cell(2, 2)
    retail_paragraph = retail_cell.paragraphs[0]
    retail_paragraph.text = "PCLN"
    set_font(retail_paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
    center_text_in_cell(retail_cell)

    # Add "Total" text and set its font
    total_cell = table.cell(3, 2)
    total_paragraph = total_cell.paragraphs[0]
    total_paragraph.text = "Agoda"
    set_font(total_paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
    center_text_in_cell(total_cell)

    total_cell = table.cell(4, 2)
    total_paragraph = total_cell.paragraphs[0]
    total_paragraph.text = "B.com"
    set_font(total_paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
    center_text_in_cell(total_cell)


    total_cell = table.cell(5, 2)
    total_paragraph = total_cell.paragraphs[0]
    total_paragraph.text = "BHI"
    set_font(total_paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
    center_text_in_cell(total_cell)

    # table.cell(15, 2).merge(table.cell(15, 1)).merge(table.cell(15, 0))
    table.cell(16, 2).merge(table.cell(16, 1)).merge(table.cell(16, 0))

    # Populate table rows with data
    for i, row_data in enumerate(df_others.itertuples(index=False), start=5):
        for j, value in enumerate(row_data, start=3):
            cell = table.cell(i, j)
            cell.text = str(value)
            center_text_in_cell(cell)
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
                
    for row in table.rows:
        row.height = Inches(0.4)
        row.height_rule = WD_ROW_HEIGHT.EXACTLY

    for cell in table.columns[0].cells: cell.width = Inches(1.0)
    for cell in table.columns[1].cells: cell.width = Inches(0.7)
    for col_idx in range(2, len(table.columns)):
        for cell in table.columns[col_idx].cells:
            cell.width = Inches(0.8)

    format_percentage_values_other(table)
    center_text_in_cell(cell)

def create_roi_table(word_document, df):
    # Header "bands" across columns with icon + title
    logos = [
        {'start_col': 2, 'end_col': 4, 'path': '../Screenshots/weeklynote/dau.png',        'title': 'DAU'},
        {'start_col': 5, 'end_col': 7, 'path': '../Screenshots/weeklynote/conversion.png', 'title': 'Conversion'},
        {'start_col': 8, 'end_col': 10, 'path': '../Screenshots/weeklynote/roi.png',        'title': 'ROI'},
    ]

    # Build table (11x11 as in your original)
    table = word_document.add_table(rows=11, cols=11, style='Table Grid')

    # Index columns sizing
    table.columns[0].width = Inches(1)
    table.columns[1].width = Inches(1)

    # Style borders/fill
    set_table_style(table, border_color="000000", border_width=8, fill_color="FFFFFF")

    header_cells = table.rows[0].cells
    remove_cell_borders(table.cell(0, 0))
    # print(title)
    table.cell(1, 0).text = 'Channel'
    
    remove_cell_borders(table.cell(1, 0))
    table.cell(0, 0).merge(table.cell(1, 0))
    for paragraph in header_cells[0].paragraphs:
        set_font(paragraph, font_name="Montserrat", font_size=Pt(14), bold=True)
    center_bottom_text_in_cell(table.cell(1, 0)) 

    
    # Merge and fill the grouped headers with icon + label
    for band in logos:
        start_col = band['start_col']
        end_col = band['end_col']
        icon_path = band['path']
        label = band['title']

        # Merge the header cells across the band
        cell = table.cell(0, start_col)
        cell.merge(table.cell(0, end_col))

        # Insert icon + label
        paragraph = cell.paragraphs[0]
        r = paragraph.add_run()
        r.add_picture(icon_path, width=Inches(0.4), height=Inches(0.4))
        # paragraph.add_run('\n' + label)
        # set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
        center_text_in_cell(cell)

    column_titles = list(df.columns)

    # Write titles starting from col 1
    for j, title in enumerate(column_titles, start=2):
        cell = table.cell(1, j)
        cell.text = str(title)
        for p in cell.paragraphs:
            set_font(p, font_name="Montserrat", font_size=Pt(10), bold=True)
        center_text_in_cell(cell)

    icons_and_labels_col1 = [
        ('../Screenshots/weeklynote/Web Marketing.png', 'Web Marketing', 3, 4),
        ('../Screenshots/weeklynote/Parity.png',       'Shop PPC',      5, 9),
    ]
    for icon_path, label, start_row, end_row in icons_and_labels_col1:
        # Merge down the first column
        cell = table.cell(start_row, 0)
        cell.merge(table.cell(end_row, 0))
        paragraph = cell.paragraphs[0]
        r = paragraph.add_run()
        r.add_picture(icon_path, width=Inches(0.4), height=Inches(0.4))
        paragraph.add_run('\n' + label)
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False)
        center_text_in_cell(cell)

    icons_and_labels_col2 = [ 
    ('../Screenshots/weeklynote/Direct.png', 'Direct', 2, 2, 0.4, 0.4), 
    ('../Screenshots/weeklynote/SEM Core.png', 'SEM Core', 3, 3, 0.4, 0.4), 
    ('../Screenshots/weeklynote/SEM Brand.png', 'SEM Brand', 4, 4, 0.4, 0.4), 
    ('../Screenshots/weeklynote/cheapflights.png', '', 5, 5, 1.0, 0.12), 
    ('../Screenshots/weeklynote/google.png', '', 6, 6, 1.0, 0.34), 
    ('../Screenshots/weeklynote/kayak.png', '', 7, 7, 1.0, 0.19), 
    ('../Screenshots/weeklynote/skyscanner.png', '', 8, 8, 1.0, 0.18), 
    ('../Screenshots/weeklynote/other.png', 'Other', 9, 9, 0.2, 0.2), 
    ('../Screenshots/weeklynote/Affiliate.png', 'Affiliate', 10, 10, 0.4, 0.4)] 
    table.cell(0, 0).merge(table.cell(1, 1))
    table.cell(2,0).merge(table.cell(2, 1))
    table.cell(10,0).merge(table.cell(10, 1))
    for icon_path, label, start_row, end_row, width, height in icons_and_labels_col2:
        cell = table.cell(start_row, 1)
        paragraph = cell.paragraphs[0] 
        paragraph.add_run().add_picture(icon_path, width=Inches(width), height=Inches(height)) 
        if label is not None and str(label).strip() != '':
            paragraph.add_run('\n' + str(label))
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False) 
        center_text_in_cell(cell)


    # Populate table rows with data
    for i, row_data in enumerate(df.itertuples(index=False), start=2):
        for j, value in enumerate(row_data, start=2):
            cell = table.cell(i, j)
            # print(i,row_data,value)
            cell.text = str(value)
            center_text_in_cell(cell)
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
           

    for row in table.rows:
        row.height = Inches(0.4)
        row.height_rule = WD_ROW_HEIGHT.EXACTLY
    for cell in table.columns[0].cells: cell.width = Inches(1.0)
    for cell in table.columns[1].cells: cell.width = Inches(1.0)
    for col_idx in range(2, len(table.columns)):
        for cell in table.columns[col_idx].cells:
            cell.width = Inches(0.8)

    format_percentage_values(table)  
    center_text_in_cell(cell) 
    set_table_width(table)

    return table
