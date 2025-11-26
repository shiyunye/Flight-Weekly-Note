from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches

def set_table_style(table, border_color="000000", border_width=4, fill_color="FFFFFF"):

    table_properties = table._element.find(qn("w:tblPr"))
    if table_properties is None:
        table_properties = OxmlElement("w:tblPr")
        table._element.insert(0, table_properties)

    # Set table borders
    tbl_borders = OxmlElement("w:tblBorders")
    for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_element = OxmlElement(f"w:{border}")
        border_element.set(qn("w:val"), "single")
        border_element.set(qn("w:sz"), str(border_width))  # Border width in eighths of a point
        border_element.set(qn("w:space"), "0")
        border_element.set(qn("w:color"), border_color)
        tbl_borders.append(border_element)
    table_properties.append(tbl_borders)

    # Set cell background color
    for row in table.rows:
        for cell in row.cells:
            # Retrieve or create the <tcPr> (table cell properties)
            cell_properties = cell._element.find(qn("w:tcPr"))
            if cell_properties is None:
                cell_properties = OxmlElement("w:tcPr")
                cell._element.insert(0, cell_properties)

            # Set cell shading (background color)
            cell_shading = OxmlElement("w:shd")
            cell_shading.set(qn("w:val"), "clear")
            cell_shading.set(qn("w:fill"), fill_color)  # Background color
            cell_properties.append(cell_shading)
            
def shade_cell(cell, color):
    """Apply background color to a cell."""
    cell_properties = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell_properties.append(shading)

def set_table_width(table, width_pct='2500'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), width_pct)
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

def adjust_cell_dimensions(table):
    if table is None:
        raise ValueError("The table is None. Make sure you created it correctly before adjusting dimensions.")
    
    table.allow_autofit = False  # Disable autofit to manually control dimensions

    for row in table.rows:
        row.height = Inches(0.4)  # Set row height

        for cell in row.cells:
            cell.width = Inches(0.8)  # Set cell width

def index_width_dimensions(table):
    """
    Adjusts the dimensions of a table, including row height and column widths.
    """
    # Disable autofit to allow manual size adjustment
    table.allow_autofit = False

    # Set uniform row height
    for row in table.rows:
        row.height = Inches(0.4)

    # Adjust column width based on content length
    for col_idx, column in enumerate(table.columns):
        max_text_length = max(
            (len(cell.text.strip()) for cell in column.cells), default=0
        )  # Get the max text length in the column
        # Estimate column width based on the longest text (adjust scaling factor if needed)
        column_width = max(Inches(1), Inches(max_text_length * 0.1))
        column.width = column_width

def center_text_in_cell(cell):
    """Center text horizontally and vertically in a table cell."""
    # Horizontal alignment
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Vertical alignment
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def center_bottom_text_in_cell(cell):
    """Center text horizontally and vertically in a table cell."""
    # Horizontal alignment
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Vertical alignment
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'bottom')
    tcPr.append(vAlign)

def set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=False, color=None):

    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
     
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)

def bold_highlight_total_cells(table):
    rows_to_highlight = set()
    cols_to_highlight = set()

    # Check the header row for 'Total' in column headers
    header_row = table.rows[0]
    for col_idx, cell in enumerate(header_row.cells):
        if "total" in cell.text.lower().strip():
            cols_to_highlight.add(col_idx)

    # Check all rows for 'Total' in the index (first cell of each row)
    for row_idx, row in enumerate(table.rows):
        if "total" in row.cells[0].text.lower().strip():
            rows_to_highlight.add(row_idx)

    # Apply formatting to the identified rows and columns
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if row_idx in rows_to_highlight or col_idx in cols_to_highlight:
                shade_cell(cell, 'D3D3D3')  # Custom shading function
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True  # Bold the text

def remove_cell_borders(cell):
    # Get or add the cell properties (tcPr)
    tc_pr = cell._element.get_or_add_tcPr()

    # Set all borders (top, bottom, left, right) to white
    for border in ["top", "bottom", "left", "right"]:
        border_element = tc_pr.find(qn(f"w:{border}"))
        if border_element is None:
            border_element = OxmlElement(f"w:{border}")
            tc_pr.append(border_element)
        
        border_element.set(qn("w:val"), "nil") 
        border_element.set(qn("w:color"), "FFFFFF")     

def format_percentage_values(table):

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    value = run.text.strip()  # Get the text content

                    # Check if the value is a string that ends with a percentage symbol
                    if isinstance(value, str) and value.endswith('%'):
                        try:
                            # Convert the value to a float, removing the '%' sign
                            num_value = float(value.strip('%'))
                            
                            # Format the value with the plus sign for positive values
                            formatted_value = f"+{num_value}%" if num_value > 0 else f"{num_value}%"
                            
                            # Set the font color based on the value (green for positive, red for negative)
                            font_color = RGBColor(106, 168, 79) if num_value > 0 else RGBColor(255, 0, 0)
                            
                            # Update the text, color, and bold the value
                            run.text = formatted_value
                            run.font.color.rgb = font_color
                            run.font.bold = True  # Make the text bold
                        except ValueError:
                            pass  # In case the conversion fails, do nothing                 

def format_percentage_values_other(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    value = run.text.strip()  # Ensure text is not empty
                    if not value:
                        continue
                    
                    font_color = RGBColor(0, 0, 0)  # Default color is black

                    try:
                        # Check if the value is a percentage
                        if value.endswith('%'):
                            num = float(value.rstrip('%'))  # Convert to float after stripping %
                            run.text = f"{num:.1f}%"  # Format with one decimal place
                        else:
                            num = float(value)  # Convert to float
                            run.font.bold = True

                            if num >= 0:
                                run.text = f"+{num:.0f}"  # Add plus sign
                                font_color = RGBColor(106, 168, 79)  # Dark green
                            else: 
                                run.text = f"{num:.0f}"  # Keep negative sign
                                font_color = RGBColor(255, 0, 0)  # Red
                    except ValueError:
                        # Leave non-numerical text unchanged
                        pass

                    # Apply font styling
                    run.font.color.rgb = font_color
                    run.font.name = "Montserrat"

def add_table_headers_index(table, title, column_titles, logos):
    
    # Table title in the first column
    
    header_cells = table.rows[0].cells
    remove_cell_borders(table.cell(0, 0))

    table.cell(1, 0).text = title
    
    remove_cell_borders(table.cell(1, 0))
    table.cell(0, 0).merge(table.cell(1, 0))
    # center_text_in_cell(table.cell(1, 0))  # Center the text in the merged cell
    center_bottom_text_in_cell(table.cell(1, 0))  # Center the text in the merged cell

    for paragraph in header_cells[0].paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(14), bold=True)

 
    for cell in header_cells:
        cell.height = Inches(2) 


    # Add headerlogos with titles
    for header_logo in logos:
        logo_cell = table.cell(0, header_logo['start_col'])
        logo_cell.merge(table.cell(0, header_logo['end_col']))
        paragraph = logo_cell.paragraphs[0]

        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_text("\n\n")
        run.add_text("\n\n")

        try:
            run.add_picture(header_logo['path'], width=Inches(0.4), height=Inches(0.4))
        except FileNotFoundError:
            paragraph.text = " "
            # print(header_logo)
        paragraph.add_run(f"\n{header_logo['title']}")
        if header_logo['title']== 'Total':
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True) 
            center_text_in_cell(cell)
        else:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
            center_text_in_cell(cell)

    # Add column titles
    for idx, col in enumerate(column_titles):
        cell = table.cell(1, idx + 1)
        col=col.split('_')[0]
        cell.text = col
        # print(col)
        center_text_in_cell(cell)
        for paragraph in cell.paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True)
        index_width_dimensions(table)

def populate_table_data(table, df):

    total_column_index = None
    total_column_indices = [col_index + 1  # Adjust for the table offset (logo column is index 0)
    for col_index, col_name in enumerate(df.columns)
    if "total" in col_name.lower()]

    for col_index, col_name in enumerate(df.columns):
        if col_name.lower() == 'total':
            total_column_index = col_index + 1

    for index, row_data in df.iterrows():
        cells = table.add_row().cells

        # Add Index logo to the first column
        index_logo_path = f'../Screenshots/weeklynote/{index}.png'
        # print(index_logo_path)
        cell = cells[0]
        paragraph = cell.paragraphs[0]
    
        run = paragraph.add_run()
        
        if str(index).lower() == "total":  # Skip if the index is "Total"
            paragraph.add_run(f"{index}")
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10),bold=True)
            center_text_in_cell(cell)
        else:
            if "(" in str(index) and ")" in str(index):
                run.add_picture(index_logo_path, height=Inches(0.4))
                center_text_in_cell(cell)
            else:
                try:
                    run.add_picture(index_logo_path, height=Inches(0.4))
                    paragraph.add_run(f"\n{index}")
                    # print(index)
                    set_font(paragraph, font_name="Montserrat", font_size=Pt(7),bold=False)
                    center_text_in_cell(cell)
                
                except FileNotFoundError:
                        paragraph.text = "logo_path miss "

        center_text_in_cell(cell)


        # Populate data cells
        for col_idx, value in enumerate(row_data):
            cell = cells[col_idx + 1]
            # print(col_idx + 1)
            cell.width = Inches(5)
            # print(value)
            cell.text = str(value)
            center_text_in_cell(cell)
        
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name="Montserrat", font_size=Pt(10))

            if total_column_index and col_idx + 1 == total_column_index:
                shade_cell(cell, 'D3D3D3')

             # Highlight "Total" columns
            if col_idx + 1 in total_column_indices:
                shade_cell(cell, 'D3D3D3')

            # print("Row Indexes:", df.index.tolist())


        
        format_percentage_values(table)
                
   
        adjust_cell_dimensions(table)


        bold_highlight_total_cells(table)

def create_word_table(df, title, logos, word_document):

    # Define dynamic index column widths for different table titles
    index_column_widths = {
        "Business": Inches(1.07),
        "Carrier": Inches(1.7),
        "Channel": Inches(1.07),
        "Source": Inches(1.07),
        "ROI": Inches(1.07)
    }

    # Set a default width if the title is not specified in the dictionary
    index_column_width = index_column_widths.get(title)
    
    # Create the table
    table = word_document.add_table(rows=2, cols=len(df.columns) + 1, style='Table Grid')
    
    # Apply table styling
    set_table_style(table, border_color="000000", border_width=8, fill_color="FFFFFF")
    
    # Adjust index column width for all rows
    table.columns[0].width = index_column_width


    # Add headers
    column_titles = df.columns.tolist()

    add_table_headers_index(table, title, column_titles, logos)

    # Populate data
    populate_table_data(table, df)

        # Remove unnecessary cell borders for the Carrier table (if title is Carrier)
    if title == "Carrier":
        for cell_coords in [
            (5, 4), (5, 5), (5, 6),
            (6, 4), (6, 5), (6, 6),
            (8, 4), (8, 5), (8, 6),
        ]:
            remove_cell_borders(table.cell(*cell_coords))

    # Remove unnecessary cell borders for the Source table (if title is Source)
    if title == "Source":
        for cell_coords in [
            (4, 1), (4, 2), (4, 3),
        ]:
            remove_cell_borders(table.cell(*cell_coords))
    
    if title == "ROI":
        for i in range(2, 10):  # i goes from 2 to 9
            for j in range(1, 4):  # j goes from 1 to 3
                remove_cell_borders(table.cell(i, j))

            

    bold_highlight_total_cells(table)

    # set_table_width(table)

def clear_document(word_document):
    """Clear all content from a Word document."""
   
    for paragraph in word_document.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Remove all tables
    for table in word_document.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)

def create_summary_table(word_document,df_summary):

    # Add the index logo to the first column
    index_logo_path = '../Screenshots/weeklynote/priceline.png'

    # Create the table
    rows, cols = len(df_summary.index) + 2, len(df_summary.columns) + 2
    table = word_document.add_table(rows=rows, cols=cols, style='Table Grid')
    
    set_table_style(table, border_color="000000", border_width=8, fill_color="FFFFFF")
    
    for cell_coords in [
        (0, 0), (6, 5),(6, 6), (6, 7), (6, 8), (6, 9),(7, 5),(7, 6), (7, 7), (7, 8), (7, 9)
    ]:
        remove_cell_borders(table.cell(*cell_coords))

    # Add the logo to the first cell and merge header cells
    header_cell = table.cell(0, 0)
    header_paragraph = header_cell.paragraphs[0]
    header_paragraph.add_run().add_picture(index_logo_path, width=Inches(1.5), height=Inches(.52))
    center_text_in_cell(header_cell)

    table.cell(0, 2).merge(table.cell(0, 1)).merge(table.cell(0, 0))
    table.cell(6,5).merge(table.cell(7, 9))


    # Define and populate column headers name
    columns = ['Metric', 'Net/Gross', 'Actual', 'Reporting Week', 'Previous Week','Reporting Week', 'Previous Week', 'MTD', 'QTD', 'YTD']
    for j, name in enumerate(columns):
        cell = table.cell(1, j)
        cell.text = name
        center_text_in_cell(cell)
        for paragraph in cell.paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True)


    # Add icons and labels for index col
    icons_and_labels = [
        ('../Screenshots/weeklynote/Tickets.png', 'Tickets', 2, 3),
        ('../Screenshots/weeklynote/Revenue.png', 'Revenue', 4, 5),
        ('../Screenshots/weeklynote/Normalized Tickets.png', 'Normalized Tickets', 6, 7)
    ]

    for icon_path, label, start_row, end_row in icons_and_labels:
        table.cell(end_row, 0).merge(table.cell(start_row, 0))
        cell = table.cell(start_row, 0)
        paragraph = cell.paragraphs[0]
        paragraph.add_run().add_picture(icon_path, width=Inches(0.4), height=Inches(0.4))
        paragraph.add_run('\n')
        paragraph.add_run(label)
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False)
        center_text_in_cell(cell)



    # Add "vs. Last Year" and "vs. Plan" labels in the header row
    last_year_logo_path = '../Screenshots/weeklynote/lastyear.png'
    plan_logo_path = '../Screenshots/weeklynote/plan.png'
    
    # "vs. Last Year"
    cell = table.cell(0, 4).merge(table.cell(0, 3))
    paragraph = cell.paragraphs[0]
    paragraph.add_run().add_picture(last_year_logo_path, width=Inches(0.4), height=Inches(0.4))
    paragraph.add_run("\n vs. Last Year")
    set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
    center_text_in_cell(cell)

    # "vs. Plan"
    cell = table.cell(0, 9).merge(table.cell(0, 5))
    paragraph = cell.paragraphs[0]
    paragraph.add_run().add_picture(plan_logo_path, width=Inches(0.4), height=Inches(0.4))
    paragraph.add_run("\n vs. Plan")
    set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
    center_text_in_cell(cell)

    # Add "Net" and "Gross" alternately in col 2 for names
    for i in range(2, len(df_summary) + 2):
        cell = table.cell(i, 1)
        cell.text = "Net" if (i - 2) % 2 == 0 else "Gross"
        paragraph = cell.paragraphs[0]  
        set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
        center_text_in_cell(cell)


    # Insert data from df_summary starting at cell (2,2)
    for i, row_data in enumerate(df_summary.itertuples(index=False), start=2):
        for j, value in enumerate(row_data, start=2):
            cell = table.cell(i, j)
            if (i, j) == (5, 2) or (i, j) == (4, 2):  
                cell.text = f"${value}"  # Add dollar sign in front of the value
            else:
                cell.text = str(value)
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name="Montserrat", font_size=Pt(10))
            center_text_in_cell(cell)
            format_percentage_values(table)

    set_table_width(table)

    adjust_cell_dimensions(table)
   

    # Adjust column widths directly within this function
    column_widths = [2, 1, 1, 1, 1, 1, 1, 1,1,1]  
    for i, width in enumerate(column_widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)

def create_others_table(word_document,df_others):
    
    # Define table dimensions
    rows, cols = (12, 9)
    table = word_document.add_table(rows=rows, cols=cols, style='Table Grid')
    set_table_style(table, border_color="000000", border_width=8, fill_color="FFFFFF")


    # Set index column widths
    table.columns[0].width = Inches(1.1)  # Adjust width for the first index column
    table.columns[1].width= Inches(0.7)

        
    # Add header logos
    header_logo_path_1 = '../Screenshots/weeklynote/priceline.png'
    header_logo_path_2 = '../Screenshots/weeklynote/expedia.png'

    # Remove unnecessary cell borders
    for cell_coords in [
        (0, 0), (1, 0), (3, 8),(4, 8), (7, 8), (8, 8), (9, 8), (10, 8), (11, 8), 
        (3, 7),(4, 7), (7, 7), (8, 7), (9, 7), (10, 7), (11, 7),
        (3, 6),(4, 6), (7, 6), (8, 6), (9, 6), (10, 6), (11, 6)
    ]:
        remove_cell_borders(table.cell(*cell_coords))

    # Merge header cells and add logos
    table.cell(0, 0).merge(table.cell(1, 2))
    header_cell_1 = table.cell(0, 3).merge(table.cell(0, 5))
    # print(header_logo_path_1)
    logo_paragraph_1 = header_cell_1.paragraphs[0]
    logo_paragraph_1.add_run().add_picture(header_logo_path_1, width=Inches(1.5), height=Inches(0.52))
    center_text_in_cell(header_cell_1)

    header_cell_2 = table.cell(0, 6).merge(table.cell(0, 8))
    logo_paragraph_2 = header_cell_2.paragraphs[0]
    logo_paragraph_2.add_run().add_picture(header_logo_path_2, width=Inches(1.4), height=Inches(0.28))
    center_text_in_cell(header_cell_2)

    # Define and populate column headers
    columns = ['Actual', 'YoY (bps)', 'YoY PW (bps)', 'Actual', 'YoY (bps)', 'YoY PW (bps)']
    for j, name in enumerate(columns):
        cell = table.cell(1, j + 3)
        cell.text = name
        center_text_in_cell(cell)
        for paragraph in cell.paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True)


    # Add icons and labels for index columns
    icons_and_labels_col1 = [
        ('../Screenshots/weeklynote/Market Share.png', 'Market Share', 2, 3),
        ('../Screenshots/weeklynote/Parity.png', 'Parity', 4, 6),
        ('../Screenshots/weeklynote/Share of Business.png', 'Share of Business', 7, 10),
    ]
    for icon_path, label, start_row, end_row in icons_and_labels_col1:
        table.cell(start_row, 0).merge(table.cell(end_row, 0))
        cell = table.cell(start_row, 0)
        paragraph = cell.paragraphs[0]
        paragraph.add_run().add_picture(icon_path, width=Inches(0.4), height=Inches(0.4))
        paragraph.add_run(label)
        paragraph.add_run('\n')
        paragraph.add_run('\n')
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False)
        center_text_in_cell(cell)
  
    icons_and_labels_col2 = [
        ('../Screenshots/weeklynote/arc.png', '', 2, 2, 0.6, 0.25),
        ('../Screenshots/weeklynote/tsa.png', '', 3, 3, 1, 0.33),
        ('../Screenshots/weeklynote/expedia.png', 'Direct vs Expedia', 4, 4, 1, 0.2),
        ('../Screenshots/weeklynote/Kayak.png', 'Kayak Placement', 5, 5, 1.2, 0.23),
        ('../Screenshots/weeklynote/Skyscanner.png', 'Skyscanner Placement', 6, 6, 1.2, 0.21),
        ('../Screenshots/weeklynote/usoutbound.png', 'US Outbound', 7, 7, 0.4, 0.4),
        ('../Screenshots/weeklynote/Merchant.png', 'Merchant', 8, 9, 0.4, 0.4),
        ('../Screenshots/weeklynote/Deals.png', 'Deals', 10, 10, 0.4, 0.4),
        ('../Screenshots/weeklynote/Bookability.png', 'Bookability', 11, 11, 0.4, 0.4),
    ]
    for icon_path, label, start_row, end_row, width, height in icons_and_labels_col2:
        if label=='Merchant':
            table.cell(start_row, 1).merge(table.cell(end_row, 1))
        else:
            table.cell(start_row, 1).merge(table.cell(end_row, 2))

        cell = table.cell(start_row, 1)
        paragraph = cell.paragraphs[0]
        
        if label in {"arc", "tsa"}:
            paragraph.add_run('\n')
            paragraph.add_run().add_picture(icon_path, width=Inches(width), height=Inches(height))
        else:
            paragraph.add_run().add_picture(icon_path, width=Inches(width), height=Inches(height))
            paragraph.add_run('\n'+label)
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=False)
        center_text_in_cell(cell)

   # Add "Retail" text and set its font
    retail_cell = table.cell(8, 2)
    retail_paragraph = retail_cell.paragraphs[0]
    retail_paragraph.text = "Retail"
    set_font(retail_paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
    center_text_in_cell(retail_cell)

    # Add "Total" text and set its font
    total_cell = table.cell(9, 2)
    total_paragraph = total_cell.paragraphs[0]
    total_paragraph.text = "Total"
    set_font(total_paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
    center_text_in_cell(total_cell)

    table.cell(3,1).merge(table.cell(3, 2))
    table.cell(11,0).merge(table.cell(11, 2))
    table.cell(3, 6).merge(table.cell(4, 8))
    table.cell(7, 6).merge(table.cell(11, 8))

    # Populate table rows with data
    for i, row_data in enumerate(df_others.itertuples(index=False), start=2):
        for j, value in enumerate(row_data, start=3):
            cell = table.cell(i, j)
            cell.text = str(value)
            center_text_in_cell(cell)
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=False)
                
    adjust_cell_dimensions(table)
    format_percentage_values_other(table)
    center_text_in_cell(cell)

def create_dau_table(word_document,df):

    logos = [
        {'start_col': 1, 'end_col': 2, 'path': '../Screenshots/weeklynote/App.png', 'title': 'App'},
        {'start_col': 3, 'end_col': 4, 'path': '../Screenshots/weeklynote/MWeb Desktop.png', 'title': 'MWeb/Desktop'},
        {'start_col': 5, 'end_col': 6, 'path': '../Screenshots/weeklynote/Total.png', 'title': 'Total'}
    ]

    # Create the table
    table = word_document.add_table(rows=2, cols=len(df.columns) + 1, style='Table Grid')
    

    # Set index column widths
    table.columns[0].width = Inches(1)  # Adjust width for the first index column
    table.columns[1].width= Inches(1)

    # Apply table styling
    set_table_style(table, border_color="000000", border_width=8, fill_color="FFFFFF")
    
    # Add headers
    column_titles = df.columns.tolist()

    # Table title in the first column
    
    header_cells = table.rows[0].cells
    remove_cell_borders(table.cell(0, 0))

    table.cell(1, 0).text = 'Channel'
    
    remove_cell_borders(table.cell(1, 0))
    table.cell(0, 0).merge(table.cell(1, 0))
    center_bottom_text_in_cell(table.cell(1, 0))  # Center the text in the merged cell

    for paragraph in header_cells[0].paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(14), bold=True)
 
    for cell in header_cells:
        cell.height = Inches(2) 

    # Add headerlogos with titles
    for header_logo in logos:
        logo_cell = table.cell(0, header_logo['start_col'])
        logo_cell.merge(table.cell(0, header_logo['end_col']))
        paragraph = logo_cell.paragraphs[0]

        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_text("\n\n")
        run.add_text("\n\n")

        try:
            run.add_picture(header_logo['path'], width=Inches(0.4), height=Inches(0.4))
        except FileNotFoundError:
            paragraph.text = " "
            # print(header_logo)
        paragraph.add_run(f"\n{header_logo['title']}")
        if header_logo['title']== 'Total':
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True) 
            center_text_in_cell(cell)
        else:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
            center_text_in_cell(cell)

    # Add column titles
    for idx, col in enumerate(column_titles):
        cell = table.cell(1, idx + 1)
        col=col.split('_')[0]
        cell.text = col
        # print(col)
        center_text_in_cell(cell)
        for paragraph in cell.paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True)
        index_width_dimensions(table)

    # Populate data
    populate_table_data(table, df)


    bold_highlight_total_cells(table)
    set_table_width(table)