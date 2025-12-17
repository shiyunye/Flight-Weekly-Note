from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT

# --- Formatting Functions ---
def format_percentage_values(table, header_row_idx=1, keywords=("Actual", "Net Tickets", "DAU", "Conversion", "ROI")):
    """Format percentage values in a table."""
    if table is None or not getattr(table, "rows", None):
        return
    if header_row_idx < 0 or header_row_idx >= len(table.rows):
        return

    keywords_lc = tuple(k.lower() for k in keywords)
    header_cells = table.rows[header_row_idx].cells
    col_is_key = [any(k in (c.text or "").strip().lower() for k in keywords_lc) for c in header_cells]
    col_is_conversion = ["conversion" in (c.text or "").strip().lower() for c in header_cells]

    for r_idx, row in enumerate(table.rows):
        if r_idx == header_row_idx:
            continue

        for c_idx, cell in enumerate(row.cells):
            is_key_col = col_is_key[c_idx] if c_idx < len(col_is_key) else False
            is_conversion_col = col_is_conversion[c_idx] if c_idx < len(col_is_conversion) else False

            for paragraph in getattr(cell, "paragraphs", []):
                for run in getattr(paragraph, "runs", []):
                    value = (run.text or "").strip()
                    if not value.endswith("%"):
                        continue

                    raw = value[:-1].replace(",", "").strip()
                    if raw in ("", "+", "-", "+.", "-."):
                        continue

                    try:
                        num_value = float(raw)
                    except ValueError:
                        continue

                    # Format value
                    if num_value > 0:
                        formatted_value = f"{num_value}%" if is_conversion_col else f"+{num_value}%"
                    elif num_value < 0:
                        formatted_value = f"{num_value}%"
                    else:
                        formatted_value = "0%"

                    # Set font color
                    if is_key_col:
                        font_color = RGBColor(0, 0, 0)
                    else:
                        font_color = RGBColor(106, 168, 79) if num_value >= 0 else RGBColor(255, 0, 0)

                    # Apply formatting
                    run.text = formatted_value
                    run.font.color.rgb = font_color
                    run.font.bold = not is_key_col


# --- Table Header Functions ---
def add_table_headers_index(table, title, column_titles, logos):
    """Add headers and logos to a table."""
    header_cells = table.rows[0].cells
    remove_cell_borders(table.cell(0, 0))
    table.cell(1, 0).text = title
    remove_cell_borders(table.cell(1, 0))
    table.cell(0, 0).merge(table.cell(1, 0))
    center_bottom_text_in_cell(table.cell(1, 0))

    for paragraph in header_cells[0].paragraphs:
        set_font(paragraph, font_name="Montserrat", font_size=Pt(14), bold=True)

    for cell in header_cells:
        cell.height = Inches(0.4)

    # Add logos
    for header_logo in logos:
        logo_cell = table.cell(0, header_logo['start_col'])
        logo_cell.merge(table.cell(0, header_logo['end_col']))
        paragraph = logo_cell.paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        try:
            run.add_picture(header_logo['path'], width=Inches(0.4), height=Inches(0.4))
        except FileNotFoundError:
            paragraph.text = " "
        paragraph.add_run(f"\n{header_logo['title']}")
        set_font(paragraph, font_name="Montserrat", font_size=Pt(7), bold=True)
        center_text_in_cell(logo_cell)

    # Add column titles
    for idx, col in enumerate(column_titles):
        cell = table.cell(1, idx + 1)
        cell.text = col.split('_')[0]
        center_text_in_cell(cell)
        for paragraph in cell.paragraphs:
            set_font(paragraph, font_name="Montserrat", font_size=Pt(10), bold=True)


# --- Table Data Population ---
def populate_table_data(table, df):
    """Populate table with data from a DataFrame."""
    total_column_indices = [
        col_index + 1 for col_index, col_name in enumerate(df.columns) if "total" in col_name.lower()
    ]

    for index, row_data in df.iterrows():
        cells = table.add_row().cells

        # Add index logo
        index_logo_path = f'../Screenshots/weeklynote/{index}.jpg'
        cell = cells[0]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        try:
            run.add_picture(index_logo_path, height=Inches(0.4), width=Inches(0.4))
            paragraph.add_run(f"\n{index}")
        except FileNotFoundError:
            paragraph.text = "logo_path miss"
        center_text_in_cell(cell)

        # Populate data cells
        for col_idx, value in enumerate(row_data):
            cell = cells[col_idx + 1]
            cell.text = str(value)
            center_text_in_cell(cell)
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name="Montserrat", font_size=Pt(10))

            # Highlight "Total" columns
            if col_idx + 1 in total_column_indices:
                shade_cell(cell, 'D3D3D3')

        format_percentage_values(table)
        adjust_cell_dimensions(table)


# --- Table Creation ---
def create_word_table(df, title, logos, word_document):
    """Create a Word table with headers and data."""
    index_column_widths = {
        "Business": Inches(1.146),
        "Carrier": Inches(1.646),
        "Channel": Inches(1),
        "Source": Inches(1),
    }
    index_column_width = index_column_widths.get(title, Inches(1))

    table = word_document.add_table(rows=2, cols=len(df.columns) + 1, style='Table Grid')
    set_table_style(table, border_color="000000", border_width=8, fill_color="FFFFFF")
    table.columns[0].width = index_column_width

    add_table_headers_index(table, title, df.columns.tolist(), logos)
    populate_table_data(table, df)

    for row in table.rows:
        row.height = Inches(0.4)
        row.height_rule = WD_ROW_HEIGHT.EXACTLY


# --- Utility Functions ---
def clear_document(word_document):
    """Clear all content from a Word document."""
    for paragraph in word_document.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    for table in word_document.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)